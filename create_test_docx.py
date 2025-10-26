from docx import Document
from pathlib import Path

def create_test_document():
    """Creates a simple .docx file for testing purposes."""
    test_data_dir = Path("tests/test_data")
    test_data_dir.mkdir(exist_ok=True)
    doc_path = test_data_dir / "test_document.docx"

    document = Document()

    document.add_heading("Test Document Title", level=1)
    document.add_paragraph("This is a test paragraph.")

    # Add a simple table
    table = document.add_table(rows=2, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Column 1'
    hdr_cells[1].text = 'Column 2'
    row_cells = table.rows[1].cells
    row_cells[0].text = 'Cell A1'
    row_cells[1].text = 'Cell B1'

    document.save(doc_path)
    print(f"Test document saved to {doc_path}")

if __name__ == "__main__":
    create_test_document()
